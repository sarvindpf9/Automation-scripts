import io
import logging
from pathlib import Path

import anthropic
import dacite
import pdfplumber
from docx import Document

from .models import Datacenter, Deliverable, ScopeSection, SowDocument

log = logging.getLogger(__name__)

# Tool schema used to force structured JSON output from Claude.
# Claude "calls" this tool with extracted fields — no free-form text to parse.
_EXTRACTION_TOOL: dict = {
    "name": "store_sow",
    "description": "Store the structured data extracted from a Platform9 SOW document.",
    "input_schema": {
        "type": "object",
        "required": ["customer", "project_title", "prepared_by"],
        "properties": {
            "customer": {
                "type": "string",
                "description": "Organisation the SOW is prepared FOR (never Platform9)",
            },
            "project_title": {"type": "string"},
            "prepared_by": {"type": "string"},
            "datacenters": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["name", "role"],
                    "properties": {
                        "name": {"type": "string"},
                        "role": {
                            "type": "string",
                            "enum": ["canary", "production"],
                        },
                        "hypervisor_count": {"type": ["integer", "null"]},
                    },
                },
            },
            "phases": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Ordered delivery phases (e.g. Canary, Production, MaaS, Migration, Training)",
            },
            "scope_sections": {
                "type": "array",
                "items": {
                    "type": "object",
                    "required": ["id", "title"],
                    "properties": {
                        "id": {"type": "string", "description": "Section number, e.g. '3.1'"},
                        "title": {"type": "string"},
                        "activities": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Key activities or tasks described in this section",
                        },
                        "deliverables": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "required": ["id", "section_id", "title", "type", "owner"],
                                "properties": {
                                    "id": {"type": "string", "description": "e.g. D-3.1-1"},
                                    "section_id": {"type": "string"},
                                    "title": {"type": "string"},
                                    "type": {
                                        "type": "string",
                                        "enum": [
                                            "mop",
                                            "runbook",
                                            "report",
                                            "training",
                                            "configuration",
                                            "tracker",
                                        ],
                                    },
                                    "owner": {
                                        "type": "string",
                                        "enum": ["Platform9", "Customer", "Joint"],
                                    },
                                    "estimated_hours": {"type": ["number", "null"]},
                                },
                            },
                        },
                    },
                },
            },
            "integrations": {
                "type": "array",
                "items": {"type": "string"},
                "description": "External systems named as integration targets (e.g. 'NetApp AFF C80', 'Microsoft AD')",
            },
            "vm_count": {
                "type": ["integer", "null"],
                "description": "Total VMs to be migrated; null if not stated",
            },
            "hypervisor_count": {
                "type": ["integer", "null"],
                "description": "Total hypervisor nodes across production environments; null if not stated",
            },
            "migration_tool": {
                "type": ["string", "null"],
                "description": "Tool named for VM migration (e.g. vJailbreak)",
            },
            "assumptions": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Bullet points from the Assumptions section verbatim",
            },
            "out_of_scope": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Bullet points from the Out of Scope section verbatim",
            },
            "deliverables": {
                "type": "array",
                "description": "Top-level formal deliverable documents listed in a Deliverables section",
                "items": {
                    "type": "object",
                    "required": ["id", "section_id", "title", "type", "owner"],
                    "properties": {
                        "id": {"type": "string"},
                        "section_id": {"type": "string"},
                        "title": {"type": "string"},
                        "type": {
                            "type": "string",
                            "enum": [
                                "mop",
                                "runbook",
                                "report",
                                "training",
                                "configuration",
                                "tracker",
                            ],
                        },
                        "owner": {
                            "type": "string",
                            "enum": ["Platform9", "Customer", "Joint"],
                        },
                        "estimated_hours": {"type": ["number", "null"]},
                    },
                },
            },
        },
    },
}

_SYSTEM_PROMPT = """\
You are a structured data extraction engine for Platform9 Statement of Work (SOW) documents.
Platform9 Private Cloud Director (PCD) is an OpenStack Epoxy+ product with a SaaS management plane.

Extract all relevant fields from the SOW text and call the store_sow tool with the structured data.

Extraction rules:
- customer: the organisation the SOW is prepared FOR (never Platform9)
- datacenters: extract each named location with its role (canary=staging/validation, production=prod) and hypervisor_count if stated
- scope_sections: one entry per numbered sub-section under "Scope of Work"; use heading content to determine type, not just the number
- activities per section: extract the key tasks described as concise action statements
- deliverables within each section: infer type from context (mop/runbook/report/training/configuration/tracker); default owner to "Platform9" unless text states otherwise
- vm_count: total VMs to migrate as an integer (e.g. "one thousand (1,000)" → 1000)
- hypervisor_count: total production hypervisor nodes as an integer
- migration_tool: named VM migration tool (e.g. vJailbreak)
- phases: ordered delivery phases inferred from objectives or scope sequence
- integrations: list every named external product/system that requires integration
- assumptions and out_of_scope: extract bullet points verbatim from those sections if present
- If a field cannot be extracted, use null for scalars or [] for arrays — never invent values
"""


def extract_text_pdf(path: Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def extract_text_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_pdf(path)
    if suffix in (".docx", ".doc"):
        return extract_text_docx(path)
    raise ValueError(f"Unsupported file type: {suffix!r} — expected .pdf or .docx")


def _call_claude_extraction(raw_text: str) -> SowDocument:
    """Send raw SOW text to Claude and return the structured SowDocument."""
    client = anthropic.Anthropic()
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        system=_SYSTEM_PROMPT,
        tools=[_EXTRACTION_TOOL],
        tool_choice={"type": "tool", "name": "store_sow"},
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"SOW Document:\n\n{raw_text}",
                        "cache_control": {"type": "ephemeral"},
                    },
                    {
                        "type": "text",
                        "text": "Extract all fields and call store_sow with the structured data.",
                    },
                ],
            }
        ],
    )

    tool_block = next(
        (b for b in response.content if b.type == "tool_use" and b.name == "store_sow"),
        None,
    )
    if tool_block is None:
        raise RuntimeError(
            f"Claude did not return a store_sow tool call. Stop reason: {response.stop_reason}"
        )

    log.debug(
        "Cache tokens — creation: %s  read: %s",
        getattr(response.usage, "cache_creation_input_tokens", "n/a"),
        getattr(response.usage, "cache_read_input_tokens", "n/a"),
    )

    return dacite.from_dict(
        data_class=SowDocument,
        data=tool_block.input,
        config=dacite.Config(strict=False),
    )


def parse_sow(file_path: str) -> SowDocument:
    """Extract structured SowDocument from a SOW PDF or DOCX on disk."""
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"SOW file not found: {path}")

    raw_text = _extract_text(path)
    log.info("Extracted %d characters from %s", len(raw_text), path.name)
    return _call_claude_extraction(raw_text)


def parse_sow_from_text(raw_text: str) -> SowDocument:
    """Extract structured SowDocument from already-extracted SOW text."""
    log.info("Parsing SOW from %d characters of provided text", len(raw_text))
    return _call_claude_extraction(raw_text)


def parse_sow_from_bytes(filename: str, content: bytes) -> SowDocument:
    """Extract structured SowDocument from in-memory PDF/DOCX bytes (Claude Desktop upload)."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if text:
                    pages.append(text)
        raw_text = "\n\n".join(pages)
    elif suffix in (".docx", ".doc"):
        doc = Document(io.BytesIO(content))
        raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Unsupported file type: {suffix!r} — expected .pdf or .docx")

    log.info("Extracted %d characters from uploaded %s", len(raw_text), filename)
    return _call_claude_extraction(raw_text)
