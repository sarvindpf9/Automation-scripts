"""Utilities for reading PDF and DOCX files from a directory into plain text."""

from __future__ import annotations

import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

log = logging.getLogger(__name__)


def read_file(path: str | Path) -> str:
    """Return the plain-text content of a single PDF or DOCX file."""
    p = Path(path).expanduser().resolve()
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(p)
    if suffix in (".docx", ".doc"):
        return _read_docx(p)
    raise ValueError(f"Unsupported file type: {suffix!r} — expected .pdf or .docx")


def read_file_from_bytes(filename: str, content: bytes) -> str:
    """Extract text from PDF or DOCX bytes in memory — no disk write required."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text(x_tolerance=2, y_tolerance=2) for page in pdf.pages]
            return "\n\n".join(p for p in pages if p)
    if suffix in (".docx", ".doc"):
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    raise ValueError(f"Unsupported file type: {suffix!r} — expected .pdf or .docx")


def read_docs_dir(docs_dir: str | Path) -> list[dict]:
    """
    Walk *docs_dir* and extract text from every .pdf and .docx file found.

    Returns a list of dicts sorted by filename:
        [{"filename": str, "path": str, "text": str}, ...]

    Files with unsupported extensions are silently skipped with a warning.
    """
    root = Path(docs_dir).expanduser().resolve()
    if not root.exists():
        log.warning("docs_dir does not exist: %s — returning empty list", root)
        return []
    if not root.is_dir():
        raise ValueError(f"docs_dir is not a directory: {root}")

    results: list[dict] = []
    for file_path in sorted(root.rglob("*")):
        if not file_path.is_file():
            continue
        suffix = file_path.suffix.lower()
        if suffix not in (".pdf", ".docx", ".doc"):
            log.debug("Skipping unsupported file: %s", file_path.name)
            continue
        try:
            text = read_file(file_path)
            results.append(
                {
                    "filename": file_path.name,
                    "path": str(file_path),
                    "text": text,
                }
            )
            log.info("Read %d chars from %s", len(text), file_path.name)
        except Exception as exc:  # noqa: BLE001 — broad intentional for resilience
            log.warning("Failed to read %s: %s", file_path.name, exc)

    return results


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> str:
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
